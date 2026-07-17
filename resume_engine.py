import os
import io
import json
import time
from pathlib import Path

from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from pypdf import PdfReader
from docx import Document

load_dotenv()
my_api_key = os.getenv("GROQ_API_KEY")
if not my_api_key:
    raise ValueError("GROQ_API_KEY not found. Add it to your .env file locally, "
        "or to Streamlit Cloud's app secrets if deployed.")

client = Groq(api_key=my_api_key)
model = "openai/gpt-oss-120b"


# ---------------------------------------------------------------------
# File readers - works on a filesystem path AND a Streamlit
# UploadedFile object (which has .name and .getvalue())
# ---------------------------------------------------------------------

def _get_bytes_and_name(file_obj):
    if hasattr(file_obj, "getvalue"):  # Streamlit UploadedFile
        return file_obj.getvalue(), file_obj.name
    else:  # plain filesystem path
        p = Path(file_obj)
        return p.read_bytes(), p.name


def read_pdf_bytes(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def read_docx_bytes(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_any_document(file_obj):
    """file_obj can be a path (str/Path) or a Streamlit UploadedFile."""
    data, name = _get_bytes_and_name(file_obj)
    suffix = Path(name).suffix.lower()
    if suffix == ".pdf":
        return read_pdf_bytes(data)
    elif suffix == ".docx":
        return read_docx_bytes(data)
    elif suffix == ".txt":
        return data.decode("utf-8", errors="ignore")
    else:
        return None


# ---------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------

class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]


class MatchResult(BaseModel):
    score: float
    details: dict


class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []


class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    total_experience_years: float | None = None
    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


jobd_schema = JobD.model_json_schema()
resume_schema = Resume.model_json_schema()


# ---------------------------------------------------------------------
# Retry wrapper - real backoff, only triggered on actual rate-limit
# errors instead of a blanket sleep after every single call
# ---------------------------------------------------------------------

def _call_groq(messages, max_retries=4):
    delay = 3
    for attempt in range(max_retries):
        try:
            return client.chat.completions.create(
                model=model,
                messages=messages,
                response_format={"type": "json_object"},
            )
        except Exception as e:
            msg = str(e).lower()
            if "rate" in msg or "429" in msg:
                time.sleep(delay)
                delay *= 2
                continue
            raise
    raise RuntimeError("Groq API failed after multiple retries (rate limited).")


# ---------------------------------------------------------------------
# LLM call 1: JD text -> structured JobD
# ---------------------------------------------------------------------

def extract_job_description(job_description_text: str) -> JobD:
    system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""
    user_prompt = f"Analyze the following job description:\n\n{job_description_text}"
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    response = _call_groq(messages)
    job_data = json.loads(response.choices[0].message.content)
    return JobD(**job_data)


# ---------------------------------------------------------------------
# LLM call 2: resume text -> structured Resume
# ---------------------------------------------------------------------

def parse_resume(resume_text: str) -> Resume:
    system_prompt = f"""
You are an expert resume parser.

Extract information from the resume based on its meaning,
not only based on exact section headings.

Different resumes may use different headings, such as:
Experience, Professional Experience, Work History, Employment, Internships.
These may all contain relevant experience.

Skills may also appear in the skills section, work experience,
internships or projects.

Return ONLY valid JSON matching this schema:

{resume_schema}

Important rules:
1. Do not invent information.
2. If a value is not available, return null.
3. If a list has no information, return an empty list.
4. Include internships inside experiences.
5. Extract skills mentioned across the entire resume.
"""
    user_prompt = f"Parse the following resume:\n\n{resume_text}"
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    response = _call_groq(messages)
    data = json.loads(response.choices[0].message.content)
    return Resume(**data)


# ---------------------------------------------------------------------
# LLM call 3: job + resume -> match score / ATS-style verdict
# ---------------------------------------------------------------------

def final_score(job: JobD, resume: Resume) -> MatchResult:
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
You are an HR recruiter / ATS system.

Compare the candidate's resume with the job description.

JOB DESCRIPTION:
{job.model_dump_json(indent=2)}

CANDIDATE RESUME:
{resume.model_dump_json(indent=2)}

Return JSON matching this schema:

{match_schema}

The "details" field must include these keys:
- matching_skills (list of strings)
- missing_skills (list of strings)
- experience_requirement_met (true/false)
- verdict (short string, 1-2 sentences)

Keep it concise and easy to read. Do not invent skills or experience
that aren't present in the resume.
"""
    messages = [{"role": "user", "content": prompt}]
    response = _call_groq(messages)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)


def analyze_resume(job: JobD, resume_file) -> dict:
    """One resume, end to end: read -> parse -> score.
    Used by both the single-resume (student) mode and the
    bulk (HR) mode - this is the one place the logic lives."""
    resume_text = read_any_document(resume_file)
    if not resume_text:
        raise ValueError("Unsupported or unreadable file.")
    parsed_resume = parse_resume(resume_text)
    result = final_score(job, parsed_resume)
    return {
        "name": parsed_resume.name or getattr(resume_file, "name", "Unknown"),
        "email": parsed_resume.email,
        "score": result.score,
        "details": result.details,
        "resume": parsed_resume,
    }